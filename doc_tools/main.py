import os
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def modify_documents_in_folder(folder_path):
    """
    此函数用于遍历指定文件夹中的所有 .docx 文件，并按照要求修改文档样式

    :param folder_path: 包含 .docx 文件的文件夹路径
    """
    # 遍历文件夹中的所有文件
    for filename in os.listdir(folder_path):
        # 检查文件是否为 .docx 格式
        if filename.endswith('.docx'):
            # 构建文件的完整路径
            file_path = os.path.join(folder_path, filename)
            try:
                # 打开文档
                doc = Document(file_path)

                # 一、页面设置
                # 获取文档的第一个节（通常是整个文档的页面设置）
                section = doc.sections[0]
                # 设置页面的上边距为 3.7 厘米
                section.top_margin = Cm(3.7)
                # 设置页面的下边距为 3.7 厘米
                section.bottom_margin = Cm(3.7)
                # 设置页面的左边距为 2.8 厘米
                section.left_margin = Cm(2.8)
                # 设置页面的右边距为 2.8 厘米
                section.right_margin = Cm(2.8)
                # 设置页眉与页面边界的距离为 1.5 厘米
                section.header_distance = Cm(1.5)
                # 设置页脚与页面边界的距离为 1.75 厘米
                section.footer_distance = Cm(1.75)

                # 标题设置
                # 假设文档的第一个段落为标题
                if doc.paragraphs:
                    title = doc.paragraphs[0]
                    # 设置标题文本的字体名称为 '方正小标宋简体'
                    title.runs[0].font.name = '方正小标宋简体'
                    # 设置标题文本的字体大小为二号字，二号字对应 22 磅
                    title.runs[0].font.size = Pt(22)
                    # 获取标题运行对象的底层 XML 元素
                    r = title.runs[0]._element
                    # 设置中文字体为 '方正小标宋简体'，确保在不同环境下正确显示
                    r.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')

                # 正文设置
                for paragraph in doc.paragraphs[1:]:  # 从第二个段落开始作为正文处理
                    # 设置段落的首行缩进为 0.7 厘米，实现左空二字的效果
                    paragraph.paragraph_format.first_line_indent = Cm(0.7)
                    # 设置段落的行距为 28 磅
                    paragraph.paragraph_format.line_spacing = Pt(28)
                    for run in paragraph.runs:
                        # 设置正文文本的中文字体名称为 '仿宋GB2312'
                        run.font.name = '仿宋GB2312'
                        r = run._element
                        # 设置中文字体为 '仿宋GB2312'，确保中文字符显示正确
                        r.rPr.rFonts.set(qn('w:eastAsia'), '仿宋GB2312')
                        # 这里暂时默认按街道办事处文件处理，设置数字和英文用 Times New Roman 字体
                        r.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
                        r.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        # 若要区分街道党委和办事处文件，可根据实际情况添加判断逻辑
                        # 例如，可通过文件名或文档中特定标识来判断

                # 八、页码
                footer = section.footer
                # 开启首页不同和奇偶页不同的设置
                sectPr = section._sectPr
                # 开启首页不同
                firstPage = sectPr.xpath('./w:titlePg')
                if not firstPage:
                    firstPageTag = OxmlElement('w:titlePg')
                    firstPageTag.set(qn('w:val'), '1')
                    sectPr.append(firstPageTag)

                # 保存修改后的文档
                doc.save(file_path)
                print(f"文件 {filename} 已成功修改。")
            except Exception as e:
                print(f"处理文件 {filename} 时出现错误: {e}")

# 使用示例
folder_path = r'D:\working\寒寒项目\代码测试\doc_tools\docs'
modify_documents_in_folder(folder_path)
