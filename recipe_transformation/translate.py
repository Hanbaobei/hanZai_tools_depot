import pandas as pd
from jinja2 import Template
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches
from docx import Document
import os
import logging
from htmldocx import HtmlToDocx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from recipe_transformation.dataBase import LexiconDataBase

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

# 获取当前文件的路径
current_file_path = os.path.abspath(__file__)
# 获取当前文件所在文件夹的位置
current_folder_path = os.path.dirname(current_file_path)


def get_lexiconDataBase():
    lexiconDataBase = LexiconDataBase()
    return lexiconDataBase


# 获取中文食谱数据
def extract_menu_data(doc_path):
    """
    从 Word 文档中提取食谱数据。
    :param doc_path: Word 文档的路径
    :return: 包含食谱数据的字典
    """
    # 打开文档
    doc = Document(doc_path)

    # 初始化数据结构
    data = {
        "餐别 星期": ["早餐", "早点", "中餐", "体弱儿餐", "午点"],
        "星期一 Monday": [],
        "星期二 Tuesday": [],
        "星期三 Wednesday": [],
        "星期四 Thursday": [],
        "星期五 Friday": [],
        "星期六 Saturday": [],
        "星期日 Sunday": [],
    }

    # 遍历文档中的表格
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            # 跳过表头行
            if i == 0:
                continue
            # 获取每一行的数据
            cells = [cell.text.strip() for cell in row.cells]
            # 将数据添加到对应的星期中
            days = [
                ("星期一 Monday", 1),
                ("星期二 Tuesday", 2),
                ("星期三 Wednesday", 3),
                ("星期四 Thursday", 4),
                ("星期五 Friday", 5),
                ("星期六 Saturday", 6),
                ("星期日 Sunday", 7),
            ]
            for day, index in days:
                if index < len(cells):
                    data[day].append(cells[index])

    return data


# 获取中英文对照表
def load_translation_from_csv(csv_path):
    """
    从 CSV 文件中加载中英文对照翻译数据。
    :param csv_path: CSV 文件的路径
    :return: 包含中英文对照的字典
    """
    try:
        df = pd.read_csv(csv_path, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(csv_path, encoding="gbk")
    translation = dict(zip(df["中文名称"], df["英文名称"]))
    return translation


def load_translation_from_db():
    """
    从数据库中加载中英文对照翻译数据。
    :return: 包含中英文对照的字典
    """
    lexiconDataBase = get_lexiconDataBase()
    rows = lexiconDataBase.get_lexicon_all()
    translation = dict(rows)
    lexiconDataBase.close()
    return translation


# 将中文食谱转换为中英文对照
def translate_text(text, translation):
    items = text.split()
    translated_items = []
    for item in items:
        trans = translation.get(item, "")
        translated_item = f"<bold>{item}</bold> <br> {trans} <br>"
        translated_items.append(translated_item)
    return " ".join(translated_items)


def translate_data(data, translation):
    """
    将食谱数据中的中文转换为中英文对照。
    :param data: 包含食谱数据的字典
    :param translation: 包含中英文对照的字典
    :return: 转换后的食谱数据字典
    """
    translated_data = {
        key: [translate_text(item, translation) for item in value]
        for key, value in data.items()
    }
    return translated_data


def generate_html(translated_data):
    """
    将翻译后的数据生成HTML内容
    :param translated_data: 翻译后的数据字典
    :return: 生成的HTML字符串
    """
    # 确保所有列表长度一致
    max_length = max(len(value) for value in translated_data.values())
    for key in translated_data:
        while len(translated_data[key]) < max_length:
            translated_data[key].append("")

    # 创建DataFrame
    df = pd.DataFrame(translated_data)

    # 从文件中加载HTML模板
    with open(
            current_folder_path + os.sep + "template.html", "r", encoding="utf-8"
    ) as f:
        template_content = f.read()
    html_template = Template(template_content)
    html = html_template.render(headers=df.columns, rows=df.values.tolist())

    return html


def generate_docx(html):
    """
    将HTML内容转换为Word文档
    :param html: 包含HTML内容的字符串
    :param output_path: 输出Word文档的路径
    """
    # 将HTML转换为Word文档
    doc = Document()
    # 创建HtmlToDocx转换器
    new_parser = HtmlToDocx()
    # 创建HtmlToDocx转换器并添加HTML内容
    new_parser.add_html_to_document(html, doc)
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # 设置页面高度为11.69英寸
    section.page_width = Inches(8.27)  # 设置页面宽度为8.27英寸
    section.orientation = WD_ORIENTATION.LANDSCAPE  # 设置为横向
    section.left_margin = Inches(0.3937)  # 1厘米转换为英寸
    section.right_margin = Inches(0.3937)
    section.top_margin = Inches(0.3937)
    section.bottom_margin = Inches(0.3937)

    # 获取文档中的第一个表格
    table = doc.tables[0] if len(doc.tables) > 0 else None
    if table:
        # 获取表格的 tblPr 元素
        tbl = table._tbl
        # 修改部分：直接获取 tblPr 元素
        tblPr = tbl.xpath("w:tblPr")[0]

        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)

    # 定义框线样式
    def set_border(element, border_type, color="000000", sz=4, space="0"):
        """
        设置表格边框样式
        :param element: 表格元素
        :param border_type: 边框类型，如 'top', 'left', 'bottom', 'right', 'insideH', 'insideV'
        :param color: 边框颜色，默认为黑色
        :param sz: 边框粗细，默认为4
        :param space: 边框间距，默认为0
        """
        tag = "w:{}".format(border_type)
        border = OxmlElement(tag)
        border.set(qn("w:color"), color)
        border.set(qn("w:sz"), str(sz))
        border.set(qn("w:space"), str(space))
        border.set(qn("w:val"), "single")
        element.append(border)

    # 为表格添加所有边框
    for border_type in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        set_border(tblBorders, border_type)

    return doc


def generate_translate_save_docx(doc_path, output_path):
    """
    :param doc_path: 输入 Word 文档的路径
    :param output_path: 输出 Word 文档的路径
    """
    # 使用示例
    data = extract_menu_data(current_folder_path + os.sep + doc_path)
    # translation = load_translation_from_csv(
    #     current_folder_path + os.sep + "translation.csv"
    # )
    translation = load_translation_from_db()
    # 转换食谱数据
    translated_data = translate_data(data, translation)

    # 生成HTML
    html = generate_html(translated_data)

    # 生成Word文档
    doc = generate_docx(html)

    # 保存Word文档
    doc.save(current_folder_path + os.sep + "data" + os.sep + output_path)


def translate_document(doc_path):
    """
    翻译Word文档中的中文文本，并保存为新的Word文档。
    :param doc_path: 输入 Word 文档的路径
    :param output_path: 输出 Word 文档的路径
    """
    # 提取食谱数据
    data = extract_menu_data(doc_path)
    # 加载中英文对照表
    translation = load_translation_from_csv(
        current_folder_path + os.sep + "translation.csv"
    )
    # 转换食谱数据
    translated_data = translate_data(data, translation)
    # 生成HTML
    html = generate_html(translated_data)
    # 生成Word文档
    doc = generate_docx(html)
    # 保存Word文档
    return doc
