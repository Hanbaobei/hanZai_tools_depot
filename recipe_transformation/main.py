import pandas as pd
from jinja2 import Template
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.shared import Inches
import os

# 获取当前文件的路径
current_file_path = os.path.abspath(__file__)
# 获取当前文件所在文件夹的位置
current_folder_path = os.path.dirname(current_file_path)


# 获取中文食谱数据
def extract_menu_data(doc_path):
    """
    从 Word 文档中提取食谱数据。
    :param doc_path: Word 文档的路径
    :return: 包含食谱数据的字典
    """
    # 打开文档
    doc = Document(doc_path)

    # 初始化数据结构
    data = {
        "餐别 星期": ["早餐", "早点", "中餐", "体弱儿餐", "午点"],
        "星期一": [],
        "星期二": [],
        "星期三": [],
        "星期四": [],
        "星期五": [],
    }

    # 遍历文档中的表格
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            # 跳过表头行
            if i == 0:
                continue
            # 获取每一行的数据
            cells = [cell.text.strip() for cell in row.cells]
            # 将数据添加到对应的星期中
            data["星期一"].append(cells[1])
            data["星期二"].append(cells[2])
            data["星期三"].append(cells[3])
            data["星期四"].append(cells[4])
            data["星期五"].append(cells[5])

    return data


# 获取中英文对照表
def load_translation_from_csv(csv_path):
    """
    从 CSV 文件中加载中英文对照翻译数据。
    :param csv_path: CSV 文件的路径
    :return: 包含中英文对照的字典
    """
    try:
        df = pd.read_csv(csv_path, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(csv_path, encoding="gbk")
    translation = dict(zip(df["中文名称"], df["英文名称"]))
    return translation


# 将中文食谱转换为中英文对照
def translate_text(text, translation):
    items = text.split()
    translated_items = [f"{item} \n {translation.get(item, '')} \n " for item in items]
    return " ".join(translated_items)

def translate_data(data, translation):
    """
    将食谱数据中的中文转换为中英文对照。
    :param data: 包含食谱数据的字典
    :param translation: 包含中英文对照的字典
    :return: 转换后的食谱数据字典
    """
    translated_data = {
        key: [translate_text(item, translation) for item in value] for key, value in data.items()
    }
    return translated_data

def translate_document(doc_path, output_path):
    """
    将 Word 文档中的中文转换为中英文对照，并保存为新的 Word 文档。
    :param doc_path: 输入 Word 文档的路径
    :param output_path: 输出 Word 文档的路径
    """
    # 使用示例
    data = extract_menu_data(current_folder_path + os.sep + doc_path)
    translation = load_translation_from_csv(
        current_folder_path + os.sep + "translation.csv"
    )
    # 转换食谱数据
    translated_data = translate_data(data, translation)

    # 创建DataFrame
    df = pd.DataFrame(translated_data)

    # 从文件中加载HTML模板
    with open(
        current_folder_path + os.sep + "template.html", "r", encoding="utf-8"
    ) as f:
        template_content = f.read()
    html_template = Template(template_content)
    html = html_template.render(headers=df.columns, rows=df.values.tolist())

    # 将HTML转换为Word文档
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE  # 设置为横向
    section.left_margin = Inches(0.3937)  # 1厘米转换为英寸
    section.right_margin = Inches(0.3937)
    section.top_margin = Inches(0.3937)
    section.bottom_margin = Inches(0.3937)

    doc.add_paragraph("尚城幼儿园幼儿食谱")
    doc.add_paragraph("周次:7")
    doc.add_paragraph("日期:2025年3月24日-28日")
    doc.add_paragraph("")

    # 添加表格
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(df.columns):
        hdr_cells[i].text = header

    for row in df.values.tolist():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell

    # 保存Word文档
    doc.save(current_folder_path + os.sep + "data" + os.sep + output_path)


translate_document("尚城幼儿园幼儿食谱第7周.docx", "尚城幼儿园幼儿食谱第7周.docx")
